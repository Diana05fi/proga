import sys
from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QWidget, QPushButton, QLabel, \
    QDialog, QComboBox, QTableWidget, QTableWidgetItem, QMessageBox, QLineEdit, QDialogButtonBox, QListWidget, \
    QListWidgetItem, QInputDialog, QFileDialog
from PyQt6.QtCore import pyqtSignal
from PyQt6.QtGui import QAction
import sqlite3
from docx import Document

# Подключение к базе данных или создание новой базы данных в памяти
conn = sqlite3.connect('example.db') # База данных SQLite в памяти
cursor = conn.cursor()

# Создание таблицы 'products' в базе данных в памяти
cursor.execute('''
    CREATE TABLE IF NOT EXISTS products (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        quantity INTEGER, 
        warehouse TEXT NOT NULL,
        delivery_point TEXT
    )
''')


class Warehouse:
    def __init__(self, name):
        self.name = name
        self.products = []
    def add_product(self, product):
        self.products.append(product)

class DeliveryPoint:
    def __init__(self, name):
        self.name = name
        self.products = []


class WindowBase(QMainWindow):
    def __init__(self):
        super().__init__()
        self.cursor = conn.cursor()

    def load_products_from_db(self):
        self.cursor.execute('SELECT * FROM products')
        products = self.cursor.fetchall()

        for row in products:
            for warehouse in self.warehouses:
                if row[3] == warehouse.name:
                    warehouse.products.append({"name": row[1], "quantity": row[2], "warehouse": row[3], "delivery_point": row[4]})
    def closeEvent(self, event):
        # Подтверждение выхода
        reply = QMessageBox.question(self, 'Message', 'Are you sure to quit?', QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            event.accept()
        else:
            event.ignore()

class UserProductDialog(QDialog):
    def __init__(self, parent=None, available_uproducts=None, delivery_points=None):
        super().__init__(parent)
        self.available_uproducts = available_uproducts if available_uproducts is not None else []
        self.delivery_points = delivery_points if delivery_points is not None else []
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout()

        name_label = QLabel("Выберите товар:")
        self.uproduct_combobox = QComboBox()
        for uproduct in self.available_uproducts:
            self.uproduct_combobox.addItem(uproduct)

        delivery_point_label = QLabel("Выберите пункт выдачи:")
        self.delivery_point_combobox = QComboBox()
        for point in self.delivery_points:
            self.delivery_point_combobox.addItem(point.name)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout.addWidget(name_label)
        layout.addWidget(self.uproduct_combobox)
        layout.addWidget(delivery_point_label)
        layout.addWidget(self.delivery_point_combobox)
        layout.addWidget(buttons)

        self.setLayout(layout)

    def accept(self):
        name = self.uproduct_combobox.currentText()
        delivery_point = self.delivery_point_combobox.currentText()

        if name and delivery_point:
            try:
                quantity, ok = QInputDialog.getInt(self, "Введите количество", f"Введите количество товара '{name}' для пункта выдачи '{delivery_point}':", 1, 1, 1000, 1)
                if ok:
                    QMessageBox.information(self, "Добавление товара", f"Товар '{name}' добавлен для пункта выдачи '{delivery_point}' в количестве {quantity}.")
                else:
                    QMessageBox.critical(self, "Ошибка", "Пожалуйста, выберите количество для товара.")
            except Exception as e:
                print(f"Произошла ошибка при добавлении товара: {e}")


class ManagerProductDialog(QDialog):
    def __init__(self, parent=None, db_connection=None, available_mproducts=None, warehouses=None, is_manager=False, mproducts=None):
        super().__init__(parent)
        self.db_connection = db_connection
        self.available_mproducts = available_mproducts if available_mproducts is not None else []
        self.warehouses = warehouses if warehouses is not None else []
        self.is_manager = is_manager
        self.mproducts = mproducts if mproducts is not None else []
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout()

        name_label = QLabel("Выберите товар:")
        self.mproduct_combobox = QComboBox()
        for mproduct in self.available_mproducts:
            self.mproduct_combobox.addItem(mproduct)

        warehouse_label = QLabel("Выберите склад:")
        self.warehouse_combobox = QComboBox()
        for warehouse in self.warehouses:
            self.warehouse_combobox.addItem(warehouse.name)

        quantity_label = QLabel("Введите количество товара:")
        self.quantity_input = QLineEdit()

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout.addWidget(name_label)
        layout.addWidget(self.mproduct_combobox)
        layout.addWidget(warehouse_label)
        layout.addWidget(self.warehouse_combobox)
        layout.addWidget(quantity_label)
        layout.addWidget(self.quantity_input)
        layout.addWidget(buttons)

        self.setLayout(layout)

    def accept(self):
        if not self.is_manager:
            QMessageBox.critical(self, "Ошибка", "У вас нет разрешения на добавление товара.")
            return

        name = self.mproduct_combobox.currentText()
        warehouse = self.warehouse_combobox.currentText()
        quantity = self.quantity_input.text()

        if name and warehouse and quantity:
            try:
                quantity = int(quantity)
                cursor = self.db_connection.cursor()
                cursor.execute(
                    'INSERT INTO products (name, quantity, warehouse) VALUES (?, ?, ?)',
                    (name, quantity, warehouse)
                )
                self.db_connection.commit()

                # Вызываем метод update_products_from_manager у родительского окна (главного окна)
                if self.parent():
                    self.parent().update_products_from_manager({
                        "name": name, "quantity": quantity, "warehouse": warehouse, "delivery_point": ""
                    })

                super().accept()
                QMessageBox.information(self, "Добавление товара", f"Товар '{name}' добавлен. Количество: {quantity}")
            except Exception as e:
                print(f"An error occurred while adding the product: {e}")
        else:
            QMessageBox.critical(self, "Ошибка", "Пожалуйста, выберите товар, склад и введите количество.")


class AuthenticationDialog(QDialog):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Авторизация")
        self.setStyleSheet("background-color: #FFC0CB;")

        layout = QVBoxLayout()

        self.username_label = QLabel("Логин:")
        self.username_input = QLineEdit()

        self.password_label = QLabel("Пароль:")
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)

        login_button = QPushButton("Войти")
        login_button.clicked.connect(self.login)

        layout.addWidget(self.username_label)
        layout.addWidget(self.username_input)
        layout.addWidget(self.password_label)
        layout.addWidget(self.password_input)
        layout.addWidget(login_button)

        self.setLayout(layout)
        self.successful = False

    def login(self):
        username = self.username_input.text()
        password = self.password_input.text()

        if username == "user" and password == "1234" or (username == "manager" and password == "password"):
            QMessageBox.information(self, "Успешно", "Авторизация прошла успешно!")
            self.successful = (True, username)  # Возвращает имя пользователя вместе с флагом успешности
            self.accept()
        else:
            QMessageBox.critical(self, "Ошибка", "Неверные учетные данные.")


class UserRestrictedMainWindow(QMainWindow):
    def __init__(self, conn, warehouses):
        super().__init__()
        self.uproducts = []
        self.selected_products = []
        self.conn = conn
        self.warehouses = warehouses
        self.history = []  # Список для хранения истории добавленных товаров
        self.manager_products = []  # Список для хранения товаров, добавленных менеджером

        self.conn = sqlite3.connect('example.db')
        self.cursor = self.conn.cursor()
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS selected_uproducts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                quantity INTEGER NOT NULL
            )
        ''')

        self.setWindowTitle("Главное окно (User)")
        self.setStyleSheet("background-color: #FFC0CB")

        self.delivery_points = [
            DeliveryPoint("Пункт выдачи 1"),
            DeliveryPoint("Пункт выдачи 2")
        ]

        uproduct_button = QPushButton("Выбрать товар", self)
        uproduct_button.clicked.connect(self.show_product_dialog)

        layout = QVBoxLayout()
        layout.addWidget(uproduct_button)

        self.uproduct_table = QTableWidget(self)
        self.uproduct_table.setColumnCount(3)
        self.uproduct_table.setHorizontalHeaderLabels(["Наименование", "Количество", "Пункт выдачи"])

        layout.addWidget(self.uproduct_table)

        central_widget = QWidget()
        central_layout = QVBoxLayout()
        central_layout.addLayout(layout)
        central_widget.setLayout(central_layout)
        self.setCentralWidget(central_widget)

    def save_selected_uproducts(self):
        try:
            for uproduct in self.uproducts:
                self.cursor.execute(
                    'INSERT INTO selected_uproducts (name, quantity) VALUES (?, ?)',
                    (uproduct['name'], uproduct['quantity'])
                )
                self.conn.commit()
            print('Выбранные товары успешно сохранены!')
        except Exception as e:
            print(f'Ошибка при сохранении выбранных товаров: {e}')

    def show_main_window(self):
        self.show()

    def show_product_dialog(self):
        product_dialog = UserProductDialog(self, available_uproducts=self.uproducts,
                                           delivery_points=self.delivery_points)
        if product_dialog.exec() == QDialog.DialogCode.Accepted:
            selected_product = product_dialog.get_selected_product()
            selected_delivery_point = product_dialog.get_selected_delivery_point()

            if selected_product and selected_delivery_point:
                self.selected_products.append({"product": selected_product, "delivery_point": selected_delivery_point})
                self.update_interface()

                self.uproducts.remove(selected_product)
                self.delivery_points.remove(selected_delivery_point)

                self.handle_remaining_products()

    def show_uproduct_dialog(self):
        uproduct_dialog = QDialog(self)
        uproduct_dialog.setWindowTitle("Добавленные товары")

        layout = QVBoxLayout()
        uproduct_table = QTableWidget()
        uproduct_table.setColumnCount(4)
        uproduct_table.setHorizontalHeaderLabels(["Наименование", "Количество", "Пункт выдачи"])

        # Здесь вы можете загрузить данные из базы данных или из других источников
        for uproduct in self.uproducts:
            rowPosition = uproduct_table.rowCount()
            uproduct_table.insertRow(rowPosition)
            uproduct_table.setItem(rowPosition, 0, QTableWidgetItem(uproduct['name']))
            uproduct_table.setItem(rowPosition, 1, QTableWidgetItem(str(uproduct['quantity'])))
            uproduct_table.setItem(rowPosition, 3, QTableWidgetItem(uproduct['delivery_point']))

        layout.addWidget(uproduct_table)
        uproduct_dialog.setLayout(layout)

        uproduct_dialog.exec()

    def closeEvent(self, event):
        reply = QMessageBox.question(self, 'Message', 'Are you sure to quit?', QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.send_close_signal()
            event.accept()
        else:
            event.ignore()

    def send_close_signal(self):
        self.save_selected_uproducts()
        self.close()
        self.show_auth_dialog()

    def show_auth_dialog(self):
        auth_dialog = AuthenticationDialog()
        if auth_dialog.exec() == QDialog.DialogCode.Accepted:
            if auth_dialog.successful[0]:
                if auth_dialog.successful[1] == "manager":
                    manager_auth_dialog = ManagerAuthenticationDialog()
                    if manager_auth_dialog.exec() == QDialog.DialogCode.Accepted:
                        manager_window = ManagerRestrictedMainWindow(self.db_connection, self.warehouses)
                        manager_window.load_products_from_db()
                        manager_window.show()
                else:
                    self.show_main_window()


    def handle_manager_products(self, selected_mproducts):
        for product in selected_mproducts:
            self.uproducts.append(product)

    def load_products_from_db(self):
        try:
            self.cursor.execute('SELECT * FROM selected_uproducts')
            products = self.cursor.fetchall()

            for row in products:
                # Обработка полученных данных
                pass
        except Exception as e:
            print(f"Error while loading products from the database: {e}")




class ManagerRestrictedMainWindow(QMainWindow):
    def __init__(self, db_connection, warehouses):
        super().__init__()
        self.db_connection = db_connection
        self.cursor = self.db_connection.cursor()
        self.warehouses = warehouses
        self.selected_products = []
        self.mproducts = []
        self.history = []
        self.close_signal = None

        # Создаем таблицу, если она не существует
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS manager_products (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                quantity INTEGER NOT NULL,
                warehouse TEXT NOT NULL
            )
        ''')

        self.setStyleSheet("background-color: #FFC0CB")
        self.setWindowTitle("Главное окно (Manager)")

        # Инициализация переменных класса
        self.warehouses = warehouses if warehouses is not None else [
            Warehouse("Склад 1"),
            Warehouse("Склад 2")
        ]
        self.selected_products = []
        self.mproducts = []
        self.history = []

        # Устанавливаем меню и виджеты
        self.setup_menu()
        self.setup_widgets()

    def setup_menu(self):
        # Создаем меню
        menu = self.menuBar().addMenu("Меню")

        # Добавляем действия в меню
        add_mproduct_action = QAction("Добавить товар", self)
        add_mproduct_action.triggered.connect(self.show_mproduct_dialog)
        menu.addAction(add_mproduct_action)

        generate_report_action = QAction("Создать отчет Word", self)
        generate_report_action.triggered.connect(self.generate_report)
        menu.addAction(generate_report_action)

        generate_plain_report_action = QAction("Создать обычный отчет", self)
        generate_plain_report_action.triggered.connect(self.generate_plain_report)
        menu.addAction(generate_plain_report_action)

    def setup_widgets(self):
        # Создаем виджеты для управления товарами
        add_mproduct_button = QPushButton("Добавить товар", self)
        add_mproduct_button.clicked.connect(self.show_mproduct_dialog)

        select_mproducts_button = QPushButton("Выбрать товары", self)
        select_mproducts_button.clicked.connect(self.select_mproducts)

        check_quantity_button = QPushButton("Проверить количество товаров", self)
        check_quantity_button.clicked.connect(self.check_mproduct_quantity)

        report_button = QPushButton("Создать отчет", self)
        report_button.clicked.connect(self.generate_plain_report)

        manage_mproducts_button = QPushButton("Управление товарами")
        manage_mproducts_button.clicked.connect(self.manage_mproducts)

        # Создаем таблицу для отображения товаров
        self.mproduct_table = QTableWidget(self)
        self.mproduct_table.setColumnCount(2)
        self.mproduct_table.setHorizontalHeaderLabels(["Наименование", "Количество"])
        self.mproduct_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)

        # Создаем макет для управляющих элементов
        layout = QVBoxLayout()
        layout.addWidget(add_mproduct_button)
        layout.addWidget(select_mproducts_button)
        layout.addWidget(check_quantity_button)
        layout.addWidget(report_button)
        layout.addWidget(manage_mproducts_button)
        layout.addWidget(self.mproduct_table)

        # Устанавливаем макет в центральный виджет
        central_widget = QWidget()
        central_layout = QVBoxLayout()
        central_layout.addLayout(layout)
        central_widget.setLayout(central_layout)
        self.setCentralWidget(central_widget)

    def select_mproducts(self):
        dialog = ManageProductsDialog(mproducts=self.mproducts, parent=self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            selected_mproducts = []
            for i in range(dialog.mproducts_list.count()):
                selected_mproducts.append({"name": dialog.mproducts_list.item(i).text()})
            self.mproducts = selected_mproducts

    def save_mproduct(self, mproduct):
        print(f"Saving mproduct: {mproduct}")
        try:
            cursor = self.db_connection.cursor()
            cursor.execute(
                'INSERT INTO products (name, warehouse, quantity) VALUES (?, ?, ?)',
                (mproduct["name"], mproduct["warehouse"], mproduct["quantity"])
            )
            self.db_connection.commit()
            print(f"Product '{mproduct['name']}' saved to the database.")
        except Exception as e:
            print(f"Error while saving product '{mproduct['name']}' to the database: {e}")

    def show_main_window(self):
        self.show()

    def closeEvent(self, event):
        self.save_selected_mproducts()
        self.show_auth_dialog()

    def update_user_products(self, product, quantity):
        product_info = {
            "name": product,
            "quantity": quantity
        }
        for user_window in self.user_windows:
            user_window.add_product_from_manager(product_info)

    def show_auth_dialog(self):
        auth_dialog = AuthenticationDialog()
        if auth_dialog.exec() == QDialog.DialogCode.Accepted:
            if auth_dialog.successful[0]:
                if auth_dialog.successful[1] == "manager":
                    manager_auth_dialog = ManagerAuthenticationDialog()
                    if manager_auth_dialog.exec() == QDialog.DialogCode.Accepted:
                        manager_window = ManagerRestrictedMainWindow(self.db_connection, self.warehouses)
                        manager_window.load_products_from_db()
                        manager_window.show()
                else:
                    self.show_main_window()

    def show_mproduct_dialog(self):
        available_mproducts = ["джинсы", "лего", "свитер"]

        # Создаем склады
        warehouse1 = Warehouse("Склад 1")
        warehouse2 = Warehouse("Склад 2")

        # Создаем диалоговое окно для добавления товара
        mproduct_dialog = ManagerProductDialog(self, db_connection=self.db_connection,
                                               available_mproducts=available_mproducts,
                                               warehouses=[warehouse1, warehouse2], is_manager=True)
        mproduct_dialog.setWindowTitle("Добавить товар")

        if mproduct_dialog.exec() == QDialog.DialogCode.Accepted:
            name = mproduct_dialog.mproduct_combobox.currentText()
            warehouse = mproduct_dialog.warehouse_combobox.currentText()
            quantity = mproduct_dialog.quantity_input.text()

            new_mproduct = {"name": name, "warehouse": warehouse, "quantity": quantity}
            self.mproducts.append(new_mproduct)

            # Сохранение товара в базу данных и обновление таблицы
            self.add_mproduct_to_db(new_mproduct)
            self.update_mproduct_table()

    def show_uproduct_dialog(self):
        uproduct_dialog = QDialog(self)
        uproduct_dialog.setWindowTitle("Просмотр товаров пользователя")

        layout = QVBoxLayout()
        uproduct_table = QTableWidget()
        uproduct_table.setColumnCount(3)
        uproduct_table.setHorizontalHeaderLabels(["Наименование", "Количество",  "Пункт выдачи"])

        for uproduct in self.uproducts:
            rowPosition = uproduct_table.rowCount()
            uproduct_table.insertRow(rowPosition)
            uproduct_table.setItem(rowPosition, 0, QTableWidgetItem(uproduct['name']))
            uproduct_table.setItem(rowPosition, 1, QTableWidgetItem(str(uproduct['quantity'])))
            uproduct_table.setItem(rowPosition, 2, QTableWidgetItem(uproduct['delivery_point']))

        layout.addWidget(uproduct_table)
        uproduct_dialog.setLayout(layout)

        uproduct_dialog.exec()

    def manage_mproducts(self):
        mproducts_dialog = ManageProductsDialog(self.mproducts)
        if mproducts_dialog.exec() == QDialog.DialogCode.Accepted:
            selected_mproducts = self.mproducts


    def check_mproduct_quantity(self):
        total_mproducts = sum(int(mproduct.get('quantity', 0)) for mproduct in self.mproducts)
        return total_mproducts

    def update_mproduct_table(self):
        self.mproduct_table.setRowCount(len(self.mproducts))

        for row, mproduct in enumerate(self.mproducts):
            self.mproduct_table.setItem(row, 0, QTableWidgetItem(mproduct['name']))
            self.mproduct_table.setItem(row, 1, QTableWidgetItem(str(mproduct['quantity'])))
        pass

    def generate_report(self):
        report_text = "Отчет по товарам:\n"
        for mproduct in self.mproducts:
            quantity = mproduct.get('quantity', 0)
            report_text += f"Товар: '{mproduct.get('name', '')}', Количество: {mproduct.get('quantity', 0)}\n"

        # Создание нового документа в формате MS Word
        doc = Document()
        doc.add_heading('Отчет Word', level=1)

        for mproduct in self.mproducts:
            p = doc.add_paragraph(f"Товар: {mproduct.get('name')}, Количество: {mproduct.get('quantity', 0)}")

        report_file, _ = QFileDialog.getSaveFileName(self, "Сохранить отчет", "", "Word files (*.docx);;All files (*.*)")
        if report_file:
            doc.save(report_file)
        print("Метод generate_report вызван")


    def generate_plain_report(self):
        report_text = "Отчет по товарам:\n"
        for mproduct in self.mproducts:
            quantity = mproduct.get('quantity', 0)
            report_text += f"Товар: '{mproduct.get('name', '')}', Количество: {mproduct.get('quantity', 0)}, Склад: {mproduct.get('warehouse', '')}\n"
        QMessageBox.information(self, "Генерация отчета", report_text)

    def save_selected_mproducts(self):
        try:
            for mproduct in self.mproducts:
                cursor = self.db_connection.cursor()
                cursor.execute(
                    'INSERT INTO manager_products (name, quantity, warehouse) VALUES (?, ?, ?)',
                    (mproduct['name'], mproduct['quantity'], mproduct['warehouse'])
                )
                self.db_connection.commit()
            print('Выбранные товары успешно сохранены!')
        except Exception as e:
            print(f'Ошибка при сохранении выбранных товаров: {e}')


    def send_close_signal(self):
        self.save_selected_mproducts()
        self.close()
        self.show_auth_dialog()

    def load_products_from_db(self):
        cursor = self.db_connection.cursor()
        print("Loading products from the database for the manager window")


    def save_selected_uproducts(self):
        try:
            # Открываем подключение к базе данных
            conn = sqlite3.connect(
                'database.db')  # Замените 'database.db' на ваш файл базы данных или используйте существующее подключение к базе данных

            # Создаем объект курсора для выполнения операций с базой данных
            cursor = conn.cursor()

            # Проходим по каждому выбранному товару и сохраняем его в базу данных
            for uproduct in self.uproducts:
                cursor.execute(
                    'INSERT INTO selected_uproducts (name, price, quantity) VALUES (?, ?, ?)',
                    (uproduct['name'], uproduct['price'], uproduct['quantity'])
                )
                # Сохраняем изменения
                conn.commit()

            # Закрываем соединение с базой данных
            conn.close()

            print('Выбранные товары успешно сохранены!')

        except Exception as e:
            print(f'Ошибка при сохранении выбранных товаров: {e}')


    def update_products_from_manager(self, manager_product):
        self.selected_products.append(manager_product)
        self.update_mproduct_table()

    def add_mproduct_to_db(self, mproduct):
        try:
            cursor.execute(
                'INSERT INTO products (name, quantity, warehouse) VALUES (?, ?, ?)',
                (mproduct['name'], mproduct['quantity'], mproduct['warehouse'])
            )
            conn.commit()
        except Exception as e:
            print(f"Error while adding product to the database: {e}")

    def send_selected_products_to_users(self, selected_products):
        for user_window in self.user_windows:
            user_window.handle_manager_products(selected_products)

    def handle_manager_products(self, selected_mproducts):
        for product in selected_mproducts:
            self.products.append(product)

        self.send_selected_products_to_users(selected_mproducts)


class ManageProductsDialog(QDialog):
    def __init__(self, mproducts, parent=None):
        super().__init__(parent)

        self.setWindowTitle("Управление товарами")
        layout = QVBoxLayout()

        self.mproducts_list = QListWidget()
        for mproduct in mproducts:
            item_text = f"{mproduct['name']} - {mproduct.get('quantity', 'Unknown')} шт."
            item = QListWidgetItem(item_text)
            self.mproducts_list.addItem(item)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok)
        buttons.accepted.connect(self.accept)

        layout.addWidget(self.mproducts_list)
        layout.addWidget(buttons)

        self.setLayout(layout)


class MainWindow(WindowBase):
    def __init__(self, db_connection):
        super().__init__()
        self.setStyleSheet("background-color: #FFC0CB")
        self.setWindowTitle("Главное окно")
        self.products = []
        self.warehouses = [
            Warehouse("Склад 1"),
            Warehouse("Склад 2")
        ]
        self.delivery_points = [
            DeliveryPoint("Пункт выдачи 1"),
            DeliveryPoint("Пункт выдачи 2")
        ]
        self.load_products_from_db(db_connection)

        product_button = QPushButton("Товары", self)
        product_button.clicked.connect(self.show_product_dialog)

        warehouse_buttons = []
        for warehouse in self.warehouses:
            button = QPushButton(warehouse.name, self)
            button.clicked.connect(lambda _, w=warehouse: self.show_warehouse_dialog(w))
            warehouse_buttons.append(button)

            view_button = QPushButton(f"Просмотр склада {warehouse.name}", self)
            view_button.clicked.connect(lambda _, w=warehouse: self.show_warehouse_view_dialog(w))
            warehouse_buttons.append(view_button)

        delivery_point_buttons = []
        for delivery_point in self.delivery_points:
            button = QPushButton(delivery_point.name, self)
            button.clicked.connect(lambda _, p=delivery_point: self.show_delivery_point_dialog(p))
            delivery_point_buttons.append(button)

        report_button = QPushButton("Отчет", self)
        report_button.clicked.connect(self.generate_report)

        layout = QVBoxLayout()


        for button in warehouse_buttons:
            layout.addWidget(button)

        for button in delivery_point_buttons:
            layout.addWidget(button)

        layout.addWidget(product_button)

        self.product_table = QTableWidget(self)
        self.product_table.setColumnCount(4)
        self.product_table.setHorizontalHeaderLabels(["Наименование", "Количество", "Склад", "Пункт выдачи"])
        self.product_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        layout.addWidget(self.product_table)

        layout.addWidget(report_button)

        central_widget = QWidget()
        central_layout = QVBoxLayout()
        central_layout.addLayout(layout)
        central_widget.setLayout(central_layout)

        self.setCentralWidget(central_widget)

    def show_product_dialog(self):
        dialog = ManagerProductDialog(self, db_connection=db_connection)  # Поправил аргумент передаваемый в диалог
        dialog.setWindowTitle("Добавление товара")
        dialog.setModal(True)
        dialog.exec()

    def update_product_table(self):
        self.product_table.setRowCount(len(self.products))
        for row, product in enumerate(self.products):
            self.product_table.setItem(row, 0, QTableWidgetItem(product['name']))
            self.product_table.setItem(row, 1, QTableWidgetItem(str(product['quantity'])))
            self.product_table.setItem(row, 2, QTableWidgetItem(product['warehouse']))
            self.product_table.setItem(row, 3, QTableWidgetItem(product['delivery_point']))



class DeliveryPointDialog(QDialog):
    product_taken = pyqtSignal(dict)

    def __init__(self, parent=None, delivery_point=None):
        super().__init__(parent)

        self.setWindowTitle(f"Пункт выдачи: {delivery_point.name}")
        self.delivery_point = delivery_point

        layout = QVBoxLayout()

        self.product_label = QLabel("Товар:")
        self.product_combobox = QComboBox()

        self.update_product_combobox()

        self.take_button = QPushButton("Взять")
        self.take_button.clicked.connect(self.take_product)

        self.add_button = QPushButton("Добавить в пункт выдачи")
        self.add_button.clicked.connect(self.add_product_to_delivery_point)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        self.buttons.rejected.connect(self.reject)

        layout.addWidget(self.product_label)
        layout.addWidget(self.product_combobox)
        layout.addWidget(self.take_button)
        layout.addWidget(self.add_button)
        layout.addWidget(self.buttons)

        self.setLayout(layout)

    def save_product_to_db(self, product):
        conn = sqlite3.connect('your_database.db')  # Подставьте свою базу данных
        cursor = conn.cursor()

        cursor.execute('''
            INSERT INTO products (name, quantity, warehouse, delivery_point)
            VALUES (?, ?, ?, ?)
        ''', (product['name'], product['quantity'], product['warehouse'], product['delivery_point']))

        conn.commit()
        conn.close()

    def closeEvent(self, event):
        self.save_product_to_db(self.delivery_point)
        event.accept()

    def update_product_combobox(self):
        self.product_combobox.clear()
        for product in self.delivery_point.products:
            if product["quantity"] > 0:
                self.product_combobox.addItem(product["name"])

    def take_product(self):
        product_name = self.product_combobox.currentText()

        for product in self.delivery_point.products:
            if product["name"] == product_name and product["quantity"] > 0:
                product["quantity"] -= 1
                self.parent().update_product_table()

                self.update_product_combobox()

                self.product_taken.emit(product)
                QMessageBox.information(self, "Успех", "Товар успешно взят.")
                return

        QMessageBox.critical(self, "Ошибка", "Нет доступных товаров для взятия.")

    def add_product_to_delivery_point(self):
        product_name = self.product_combobox.currentText()

        for product in self.parent().products:
            if product["name"] == product_name:
                self.delivery_point.add_product(product)
                self.update_product_combobox()
                self.parent().update_product_table()

                QMessageBox.information(self, "Успех", "Товар успешно добавлен в пункт выдачи.")
                return

        QMessageBox.critical(self, "Ошибка", "Выберите товар для добавления в пункт выдачи.")


class TakeProductDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.setWindowTitle("Получить товар из пункта выдачи")

        layout = QVBoxLayout()

        uproduct_label = QLabel("Выберите товар для получения:")
        self.uproduct_combobox = QComboBox()

        take_button = QPushButton("Получить товар")
        take_button.clicked.connect(self.take_product)

        layout.addWidget(uproduct_label)
        layout.addWidget(self.uproduct_combobox)
        layout.addWidget(take_button)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(self.reject)

        layout.addWidget(buttons)

        self.setLayout(layout)

    def take_product(self):
        uproduct_name = self.uproduct_combobox.currentText()

        for delivery_point in self.parent().delivery_points:
            for uproduct in delivery_point.products:
                if uproduct['name'] == uproduct_name and uproduct['quantity'] > 0:
                    # Получение товара
                    uproduct['quantity'] -= 1
                    self.parent().update_uproduct_table()
                    QMessageBox.information(self, "Успешно", "Товар успешно получен из пункта выдачи.")
                    return
        QMessageBox.critical(self, "Ошибка", "Выбранный товар отсутствует в пункте выдачи или закончился.")



class ManagerAuthenticationDialog(QDialog):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Авторизация для менеджера")
        self.setStyleSheet("background-color: #FFC0CB;")

        layout = QVBoxLayout()

        self.username_label = QLabel("Логин:")
        self.username_input = QLineEdit()

        self.password_label = QLabel("Пароль:")
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)

        login_button = QPushButton("Войти")
        login_button.clicked.connect(self.login)

        layout.addWidget(self.username_label)
        layout.addWidget(self.username_input)
        layout.addWidget(self.password_label)
        layout.addWidget(self.password_input)
        layout.addWidget(login_button)

        self.setLayout(layout)

    def login(self):
        username = self.username_input.text()
        password = self.password_input.text()

        if username == "manager" and password == "password":
            QMessageBox.information(self, "Успешно", "Авторизация для менеджера прошла успешно!")
            self.accept()
        else:
            QMessageBox.critical(self, "Ошибка", "Неверные учетные данные.")



class WarehouseDialog(QDialog):
    def __init__(self, parent=None, warehouse=None):
        super().__init__(parent)

        self.setWindowTitle(f"Склад: {warehouse.name}")

        layout = QVBoxLayout()

        mproduct_label = QLabel("Товар:")
        self.mproduct_combobox = QComboBox()
        for mproduct in parent.mproducts:
            if mproduct["warehouse"] == warehouse.name:
                self.mproduct_combobox.addItem(mproduct["name"])

        add_button = QPushButton("Добавить")
        add_button.clicked.connect(self.add_mproduct)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(self.reject)

        layout.addWidget(mproduct_label)
        layout.addWidget(self.mproduct_combobox)
        layout.addWidget(add_button)
        layout.addWidget(buttons)

        self.setLayout(layout)

        self.warehouse = warehouse

    def add_product(self):
        mproduct_name = self.mproduct_combobox.currentText()

        # Найти выбранный товар
        for mproduct in self.parent().mproducts:
            if mproduct["name"] == mproduct_name:

                # Добавить товар на выбранный склад
                self.warehouse.add_mproduct(mproduct)

                QMessageBox.information(self, "Успех", "Товар успешно добавлен на склад.")
                break



class WarehouseViewDialog(QDialog):
    def __init__(self, parent=None, warehouse=None):
        super().__init__(parent)

        self.setWindowTitle(f"Просмотр склада: {warehouse.name}")

        layout = QVBoxLayout()

        product_list = QListWidget()
        for product in warehouse.products:
            product_list.addItem(f"{product['name']} ({product['quantity']})")

        layout.addWidget(product_list)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(self.reject)

        layout.addWidget(buttons)

        self.setLayout(layout)


class AuthenticationController:
    def __init__(self, conn, warehouses):
        self.conn = conn
        self.warehouses = warehouses
        self.auth_dialog = AuthenticationDialog()
        self.auth_dialog.accepted.connect(self.handle_auth_accepted)

    def create_initial_window(self):
        self.auth_dialog = AuthenticationDialog()
        self.auth_dialog.accepted.connect(self.handle_auth_accepted)
        self.auth_dialog.show()

        self.manager_window = ManagerRestrictedMainWindow(self.conn, self.warehouses)
        self.user_window = UserRestrictedMainWindow(self.conn, self.warehouses)


        self.manager_window.exit_button = QPushButton("Выход")
        self.manager_window.exit_button.clicked.connect(
            self.manager_window.show_auth_dialog)  # Кнопка "Выход" для менеджера вызывает окно аутентификации

        self.user_window.exit_button = QPushButton("Выход")
        self.user_window.exit_button.clicked.connect(
            self.user_window.show_auth_dialog)  # Кнопка "Выход" для пользователя вызывает окно аутентификации

    def handle_auth_accepted(self):
        if self.auth_dialog.successful[0]:
            if self.auth_dialog.successful[1] == "manager":
                self.show_manager_window()
            else:
                self.show_user_window()

    def show_user_window(self):
        self.user_window = UserRestrictedMainWindow(self.conn, self.warehouses)
        self.user_window.db_connection = self.conn
        self.user_window.load_products_from_db()
        self.user_window.show()

    def show_manager_window(self):
        manager_auth_dialog = ManagerAuthenticationDialog()
        if manager_auth_dialog.exec() == QDialog.DialogCode.Accepted:
            self.manager_window = ManagerRestrictedMainWindow(self.conn, self.warehouses)
            self.manager_window.load_products_from_db()
            self.manager_window.show()

def main():
    app = QApplication([])

    conn = sqlite3.connect(':memory:')
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            quantity INTEGER,
            warehouse TEXT NOT NULL,
            delivery_point TEXT
        )
    ''')

    warehouses = [Warehouse("Warehouse 1"), Warehouse("Warehouse 2")]

    auth_controller = AuthenticationController(conn, warehouses)

    layout = QVBoxLayout()

    exit_button = QPushButton("Выход")
    exit_button.clicked.connect(auth_controller.create_initial_window)
    layout.addWidget(exit_button)

    main_window = QWidget()
    main_window.setLayout(layout)
    main_window.show()

    sys.exit(app.exec())

    manager_auth_dialog = ManagerAuthenticationDialog()
    if manager_auth_dialog.exec() == QDialog.DialogCode.Accepted:
        manager_window = ManagerRestrictedMainWindow(conn, warehouses)
        manager_window.load_products_from_db()
        manager_window.show()
        user_window = UserRestrictedMainWindow(conn, warehouses)
        user_window.load_products_from_db()
        user_window.show()

if __name__ == "__main__":
    main()


